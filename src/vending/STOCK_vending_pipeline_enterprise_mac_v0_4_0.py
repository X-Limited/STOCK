"""
STOCK ENTERPRISE VENDING ANALYTICS PIPELINE
Version 0.4.0
MAC STABLE + PARQUET SAFE EDITION
"""

import multiprocessing as mp
import os
import tkinter as tk
import warnings
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from tkinter import filedialog

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches

warnings.filterwarnings("ignore")
CPU_COUNT = max(mp.cpu_count() - 1, 1)


def detect_column(df, candidates):
    normalized = {str(c).strip().lower(): c for c in df.columns}
    for candidate in candidates:
        key = candidate.strip().lower()
        if key in normalized:
            return normalized[key]
    return None


def classify(fill):
    if pd.isna(fill):
        return "UNKNOWN"
    if fill == 0:
        return "EMPTY"
    if fill <= 20:
        return "CRITICAL"
    if fill <= 40:
        return "LOW"
    if fill <= 70:
        return "NORMAL"
    return "FULL"


def load_stock_file(file):
    try:
        print(f"Loading: {os.path.basename(file)}")
        df = pd.read_excel(file, engine="openpyxl")
        df.columns = [str(c).strip() for c in df.columns]

        forbidden = ["CZ automatu", "Středisko", "Název místa", "Umístění"]

        if any(col in df.columns for col in forbidden):
            print(f"SKIPPED LOCATION FILE: {os.path.basename(file)}")
            return pd.DataFrame()

        df["source_file"] = os.path.basename(file)
        return df

    except Exception as e:
        print("ERROR:", file, e)
        return pd.DataFrame()


def main():
    root = tk.Tk()
    root.withdraw()

    stock_files = filedialog.askopenfilenames(
        title="Select STOCK XLSX files", filetypes=[("Excel files", "*.xlsx")]
    )

    if len(stock_files) == 0:
        raise Exception("No stock files selected.")

    location_file = filedialog.askopenfilename(
        title="Select LOCATION XLSX file", filetypes=[("Excel files", "*.xlsx")]
    )

    if not location_file:
        raise Exception("No location file selected.")

    output_dir = Path.home() / "Desktop" / "STOCK_OUTPUT"
    output_dir.mkdir(parents=True, exist_ok=True)

    locations = pd.read_excel(location_file)
    locations.columns = [str(c).strip() for c in locations.columns]

    vm_col = detect_column(locations, ["CZ automatu", "VM", "MachineNumber", "Automat"])

    locations["vm_id"] = locations[vm_col].astype(str)

    with ThreadPoolExecutor(max_workers=CPU_COUNT) as executor:
        frames = list(executor.map(load_stock_file, stock_files))

    frames = [f for f in frames if len(f) > 0]
    stock = pd.concat(frames, ignore_index=True)

    machine_col = detect_column(
        stock, ["MachineNumber", "Machine Number", "VM", "Machine", "Automat"]
    )

    product_col = detect_column(stock, ["Product/ComponentName", "Product", "ComponentName"])

    capacity_col = detect_column(stock, ["Product/Component capacity", "Capacity"])

    fill_col = detect_column(stock, ["Product/Component Fill quantity", "Fill quantity", "Fill"])

    stock["vm_id"] = stock[machine_col].astype(str)
    stock["capacity"] = pd.to_numeric(stock[capacity_col], errors="coerce")
    stock["fill_qty"] = pd.to_numeric(stock[fill_col], errors="coerce")

    stock = stock[stock["capacity"].fillna(-1) >= 0]

    stock["fill_percent"] = np.where(
        stock["capacity"] > 0, (stock["fill_qty"] / stock["capacity"]) * 100, np.nan
    )

    stock["status"] = stock["fill_percent"].apply(classify)

    merged = stock.merge(locations, on="vm_id", how="left")

    machine_summary = (
        merged.groupby("vm_id")
        .agg(
            avg_fill=("fill_percent", "mean"),
            critical=("status", lambda x: (x == "CRITICAL").sum()),
            empty=("status", lambda x: (x == "EMPTY").sum()),
        )
        .reset_index()
    )

    machine_summary["risk_score"] = (
        machine_summary["critical"] * 2
        + machine_summary["empty"] * 3
        + (100 - machine_summary["avg_fill"].fillna(0))
    )

    top_risk = machine_summary.sort_values("risk_score", ascending=False).head(50)

    for col in merged.columns:
        if merged[col].dtype == "object":
            merged[col] = merged[col].astype(str)

    merged.to_parquet(output_dir / "merged_data.parquet", index=False, engine="pyarrow")

    merged.head(1000000).to_csv(output_dir / "merged_data_sample.csv", index=False)

    with pd.ExcelWriter(output_dir / "STOCK_VENDING_ANALYTICS.xlsx", engine="xlsxwriter") as writer:
        machine_summary.to_excel(writer, sheet_name="Machine_Summary", index=False)

        top_risk.to_excel(writer, sheet_name="Top_Risk_VM", index=False)

    plt.figure(figsize=(12, 6))
    plt.bar(top_risk["vm_id"].astype(str), top_risk["risk_score"])
    plt.xticks(rotation=90)
    plt.tight_layout()

    chart = output_dir / "risk_machines.png"
    plt.savefig(chart)
    plt.close()

    doc = Document()
    doc.add_heading("STOCK EXECUTIVE REPORT", level=1)
    doc.add_paragraph(f"Machines: {merged['vm_id'].nunique():,}")
    doc.add_paragraph(f"Rows: {len(merged):,}")
    doc.add_paragraph(f"Threads: {CPU_COUNT}")
    doc.add_picture(str(chart), width=Inches(6))
    doc.save(output_dir / "EXECUTIVE_REPORT.docx")

    print("PIPELINE FINISHED")
    print(output_dir)


if __name__ == "__main__":
    mp.freeze_support()
    main()
