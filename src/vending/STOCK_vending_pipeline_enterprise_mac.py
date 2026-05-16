"""
STOCK ENTERPRISE VENDING ANALYTICS PIPELINE
FIXED MAC VERSION

Fixes:
- avoids macOS ProcessPool BrokenProcessPool errors by using ThreadPoolExecutor
- requires stock files and location file to be selected separately
- skips location/master files accidentally selected with stock files
- robust case-insensitive column detection
- exports Excel, Parquet, chart and DOCX executive report
"""

import multiprocessing as mp
import os
import tkinter as tk
import warnings
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from tkinter import filedialog

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches

warnings.filterwarnings("ignore")
CPU_COUNT = max(mp.cpu_count() - 1, 1)


def detect_column(df, candidates):
    normalized = {str(c).strip().lower(): c for c in df.columns}
    for candidate in candidates:
        key = candidate.strip().lower()
        if key in normalized:
            return normalized[key]
    return None


def classify(fill):
    if pd.isna(fill):
        return "UNKNOWN"
    if fill == 0:
        return "EMPTY"
    if fill <= 20:
        return "CRITICAL"
    if fill <= 40:
        return "LOW"
    if fill <= 70:
        return "NORMAL"
    return "FULL"


def load_stock_file(file):
    try:
        print(f"Loading: {os.path.basename(file)}")
        df = pd.read_excel(file, engine="openpyxl")
        df.columns = [str(c).strip() for c in df.columns]

        location_markers = ["CZ automatu", "Středisko", "Název místa", "Umístění"]
        if any(col in df.columns for col in location_markers):
            print(f"SKIPPED LOCATION/MASTER FILE: {os.path.basename(file)}")
            return pd.DataFrame()

        df["source_file"] = os.path.basename(file)
        return df
    except Exception as exc:
        print("ERROR:", file, exc)
        return pd.DataFrame()


def main():
    print("\n===================================================")
    print("STOCK ENTERPRISE VENDING ANALYTICS")
    print("===================================================")
    print(f"CPU CORES AVAILABLE: {mp.cpu_count()}")
    print(f"THREADS USED: {CPU_COUNT}")
    print("===================================================")

    root = tk.Tk()
    root.withdraw()

    print("\nSELECT ONLY STOCK XLSX FILES")
    print("DO NOT SELECT LOCATION FILE HERE")
    stock_files = filedialog.askopenfilenames(
        title="Select STOCK XLSX files only", filetypes=[("Excel files", "*.xlsx")]
    )
    if len(stock_files) == 0:
        raise Exception("No stock files selected.")

    print("\nSELECT LOCATION FILE")
    location_file = filedialog.askopenfilename(
        title="Select LOCATION XLSX file", filetypes=[("Excel files", "*.xlsx")]
    )
    if not location_file:
        raise Exception("No location file selected.")

    output_dir = Path.home() / "Desktop" / "STOCK_OUTPUT"
    output_dir.mkdir(parents=True, exist_ok=True)

    print("\nLoading location file...")
    locations = pd.read_excel(location_file)
    locations.columns = [str(c).strip() for c in locations.columns]

    location_vm_col = detect_column(locations, ["CZ automatu", "VM", "MachineNumber", "Automat"])
    if location_vm_col is None:
        raise Exception(f"Location VM column not found. Columns: {list(locations.columns)}")
    locations["vm_id"] = locations[location_vm_col].astype(str)
    print(f"Locations loaded: {len(locations):,}")

    print("\nParallel loading stock files...")
    with ThreadPoolExecutor(max_workers=CPU_COUNT) as executor:
        frames = list(executor.map(load_stock_file, stock_files))
    frames = [frame for frame in frames if len(frame) > 0]
    if len(frames) == 0:
        raise Exception("No valid stock files loaded. Check that you selected Stock_*.xlsx files.")

    stock = pd.concat(frames, ignore_index=True)
    print(f"Total rows loaded: {len(stock):,}")
    print("Detected stock columns:")
    print(list(stock.columns))

    machine_col = detect_column(
        stock, ["MachineNumber", "Machine Number", "VM", "Machine", "Automat"]
    )
    product_col = detect_column(stock, ["Product/ComponentName", "Product", "ComponentName"])
    capacity_col = detect_column(stock, ["Product/Component capacity", "Capacity"])
    fill_col = detect_column(stock, ["Product/Component Fill quantity", "Fill quantity", "Fill"])

    missing = []
    if machine_col is None:
        missing.append("machine column")
    if capacity_col is None:
        missing.append("capacity column")
    if fill_col is None:
        missing.append("fill quantity column")
    if missing:
        raise Exception(
            f"Missing required columns: {missing}. Available columns: {list(stock.columns)}"
        )

    stock["vm_id"] = stock[machine_col].astype(str)
    stock["capacity"] = pd.to_numeric(stock[capacity_col], errors="coerce")
    stock["fill_qty"] = pd.to_numeric(stock[fill_col], errors="coerce")
    stock = stock[stock["capacity"].fillna(-1) >= 0]
    stock["fill_percent"] = np.where(
        stock["capacity"] > 0, (stock["fill_qty"] / stock["capacity"]) * 100, np.nan
    )
    stock["missing_qty"] = stock["capacity"] - stock["fill_qty"]
    stock["status"] = stock["fill_percent"].apply(classify)

    merged = stock.merge(locations, on="vm_id", how="left")

    agg = {
        "avg_fill": ("fill_percent", "mean"),
        "min_fill": ("fill_percent", "min"),
        "max_fill": ("fill_percent", "max"),
        "critical": ("status", lambda x: (x == "CRITICAL").sum()),
        "empty": ("status", lambda x: (x == "EMPTY").sum()),
    }
    if product_col is not None:
        agg["products"] = (product_col, "nunique")

    machine_summary = merged.groupby("vm_id").agg(**agg).reset_index()
    machine_summary["risk_score"] = (
        machine_summary["critical"] * 2
        + machine_summary["empty"] * 3
        + (100 - machine_summary["avg_fill"].fillna(0))
    )
    machine_summary = machine_summary.sort_values("risk_score", ascending=False)

    print("\nExporting outputs...")
    merged.to_parquet(output_dir / "merged_data.parquet", index=False)

    with pd.ExcelWriter(output_dir / "STOCK_VENDING_ANALYTICS.xlsx", engine="xlsxwriter") as writer:
        merged.to_excel(writer, sheet_name="Merged_Data", index=False)
        machine_summary.to_excel(writer, sheet_name="Machine_Summary", index=False)

    top_risk = machine_summary.head(15)
    risk_chart = output_dir / "risk_machines.png"
    plt.figure(figsize=(12, 6))
    plt.bar(top_risk["vm_id"].astype(str), top_risk["risk_score"])
    plt.xticks(rotation=90)
    plt.ylabel("Risk Score")
    plt.title("Top Risk Machines")
    plt.tight_layout()
    plt.savefig(risk_chart)
    plt.close()

    doc = Document()
    doc.add_heading("STOCK EXECUTIVE REPORT", level=1)
    doc.add_paragraph(f"Machines: {merged['vm_id'].nunique():,}")
    doc.add_paragraph(f"Rows: {len(merged):,}")
    doc.add_paragraph(f"Threads: {CPU_COUNT}")
    doc.add_picture(str(risk_chart), width=Inches(6))
    doc.save(output_dir / "EXECUTIVE_REPORT.docx")

    print("\n===================================================")
    print("PIPELINE FINISHED")
    print("===================================================")
    print("OUTPUT:", output_dir)


if __name__ == "__main__":
    mp.freeze_support()
    main()
